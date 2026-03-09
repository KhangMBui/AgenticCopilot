"""
Workspace management endpoints
"""

from io import BytesIO
from pathlib import Path

from app.models.document import Document
from app.settings import settings
from core.embeddings.openai_client import OpenAIEmbeddingsClient
from core.ingestion import ingest_document_chunks
from fastapi import APIRouter, Depends, HTTPException, status, File, UploadFile
from pypdf import PdfReader
from docx import Document as DocxDocument
from sqlalchemy.orm import Session
from sqlalchemy import select

from app.db import get_db
from app.models import Workspace
from app.schemas.workspaces import WorkspaceCreateRequest, WorkspaceResponse


router = APIRouter(prefix="/workspaces", tags=["workspaces"])


@router.post("", response_model=WorkspaceResponse, status_code=status.HTTP_201_CREATED)
def create_workspace(request: WorkspaceCreateRequest, db: Session = Depends(get_db)):
    """Create a new workspace."""
    workspace = Workspace(name=request.name)
    db.add(workspace)
    db.commit()
    db.refresh(workspace)
    return workspace


@router.get("", response_model=list[WorkspaceResponse])
def list_workspaces(db: Session = Depends(get_db)):
    """List all workspaces."""
    stmt = select(Workspace).order_by(Workspace.created_at.desc())
    workspaces = db.scalars(stmt).all()
    return workspaces


@router.get("/{workspace_id}", response_model=WorkspaceResponse)
def get_workspace(workspace_id: int, db: Session = Depends(get_db)):
    """Get workspace by ID."""
    workspace = db.get(Workspace, workspace_id)
    if not workspace:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Workspace {workspace_id} not found",
        )
    return workspace


@router.post("/{workspace_id}/documents/upload")
async def upload_document(
    workspace_id: int,
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
):
    try:
        workspace = db.get(Workspace, workspace_id)
        if not workspace:
            raise HTTPException(status_code=404, detail="Workspace not found")

        raw = await file.read()
        if not raw:
            raise HTTPException(status_code=400, detail="Empty file")

        content, mime = _extract_text_from_upload(file, raw)

        doc = Document(
            workspace_id=workspace_id,
            filename=file.filename or "uploaded_file",
            mime_type=mime,
            content=content,
        )
        db.add(doc)
        db.flush()

        # existing chunk + embed pipeline here
        embeddings = OpenAIEmbeddingsClient(api_key=settings.openai_api_key)
        ingest_document_chunks(db=db, document=doc, embeddings=embeddings)

        db.commit()
        db.refresh(doc)

        return {
            "document_id": doc.id,
            "filename": doc.filename,
            "mime_type": doc.mime_type,
        }
    except Exception:
        db.rollback()
        raise


def _extract_text_from_upload(file: UploadFile, raw: bytes) -> tuple[str, str]:
    """
    Returns (content_text, mime_type).
    Supports text/*, pdf, docx.
    """
    mime = file.content_type or "application/octet-stream"
    filename = (file.filename or "uploaded_file").lower()
    ext = Path(filename).suffix

    # text-like files
    if mime.startswith("text/") or ext in {
        ".py",
        ".js",
        ".ts",
        ".java",
        ".go",
        ".rs",
        ".md",
        ".json",
        ".yaml",
        ".yml",
        ".txt",
    }:
        return raw.decode("utf-8", errors="replace"), mime

    # PDF
    if mime == "application/pdf" or ext == ".pdf":
        reader = PdfReader(BytesIO(raw))
        pages = [p.extract_text() or "" for p in reader.pages]
        text = "\n".join(pages).strip()
        if not text:
            raise HTTPException(status_code=400, detail="PDF has no extractable text")
        return text, "application/pdf"

    # DOCX
    if (
        mime
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        or ext == ".docx"
    ):
        docx = DocxDocument(BytesIO(raw))
        paragraphs = [p.text for p in docx.paragraphs if p.text]
        text = "\n".join(paragraphs).strip()
        if not text:
            raise HTTPException(status_code=400, detail="DOCX has no extractable text")
        return (
            text,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    raise HTTPException(
        status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
        detail=f"Unsupported file type: {mime} ({ext})",
    )
